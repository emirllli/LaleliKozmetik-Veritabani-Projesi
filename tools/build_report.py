from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = ROOT / "diagrams"
OUTPUT = DOCS / "Laleli_Kozmetik_Final_Rapor.docx"
ER_IMAGE = DIAGRAMS / "yeni_veritabani_semasi.png"
PRODUCT_SCREEN = DIAGRAMS / "ui_screenshots" / "01_urun_yonetimi.png"
CUSTOMER_SCREEN = DIAGRAMS / "ui_screenshots" / "02_musteri_kaydi.png"
SALE_SCREEN = DIAGRAMS / "ui_screenshots" / "03_satis_ekrani.png"
CATEGORY_SCREEN = DIAGRAMS / "ui_screenshots" / "04_kategoriler.png"


def draw_er_diagram():
    pass


def draw_screen_mockup(path, title, fields, columns):
    pass


def draw_screen_images():
    pass


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level == 1 else 120)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc, title, code):
    add_heading(doc, title, 2)
    for line in code.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
        set_cell_shading(hdr[index], "D9EAF7")
        hdr[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def main():
    DOCS.mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Laleli Kozmetik ve Kişisel Bakım Mağazası\nVeritabanı Yönetim Sistemleri II Final Ödevi")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Emirhan Laleli\n23010708033\nBTS304 - 2026").bold = True
    doc.add_page_break()

    add_heading(doc, "ADIM-1: Senaryo ve Problem Tanımı")
    doc.add_paragraph(
        "Bu proje, makyaj, cilt bakımı, parfüm ve kişisel bakım ürünleri satan Laleli Kozmetik ve "
        "Kişisel Bakım Mağazası için hazırlanmıştır. Mağazada ürün giriş-çıkışları, müşteri kayıtları "
        "ve satış hareketleri tek bir veritabanı üzerinden takip edilir."
    )
    add_bullets(
        doc,
        [
            "Problem: Stok takibi defter veya dağınık dosyalarla yapıldığı için güncel stok miktarı hızlı görülememektedir.",
            "Problem: Müşterilere yapılan satışlar raporlanamadığı için müşteri bazlı toplam harcama izlenememektedir.",
            "Hedef: Satış yapıldığında ürün stok miktarını otomatik azaltan ve satış raporu sunan bir sistem geliştirmek.",
            "Kısıt: Kayıtlı olmayan müşteriye veya kayıtlı olmayan ürüne satış yapılmaz.",
            "Kısıt: Stok miktarı yetersizse satış işlemi BLL ve trigger tarafında engellenir.",
        ],
    )

    add_heading(doc, "ADIM-2: Varlıklar ve Nitelikler")
    add_bullets(
        doc,
        [
            "Kategoriler (Kategori ID, Kategori Adı, Açıklama)",
            "Ürünler (Ürün ID, Kategori ID, Ürün Adı, Marka, Birim Fiyat, Stok Miktarı, Barkod)",
            "Müşteriler (Müşteri ID, Ad, Soyad, Telefon, E-posta, Adres)",
            "Satışlar (Satış ID, Ürün ID, Müşteri ID, Adet, Birim Fiyat, Toplam Tutar, Satış Tarihi)",
        ],
    )

    add_heading(doc, "Varlıklar Arası İlişkiler", level=2)
    
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Kategori-Ürün").bold = True
    p2 = doc.add_paragraph(style="List Bullet")
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.add_run("Bir kategoride birden fazla ürün bulunabilir.")
    p3 = doc.add_paragraph(style="List Bullet")
    p3.paragraph_format.left_indent = Inches(0.4)
    p3.add_run("1:N")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Ürün-Satış").bold = True
    p2 = doc.add_paragraph(style="List Bullet")
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.add_run("Bir ürün birden fazla satışta bulunabilir.")
    p3 = doc.add_paragraph(style="List Bullet")
    p3.paragraph_format.left_indent = Inches(0.4)
    p3.add_run("1:N")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Müşteri-Satış").bold = True
    p2 = doc.add_paragraph(style="List Bullet")
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.add_run("Bir müşteri birden fazla satın alma yapabilir.")
    p3 = doc.add_paragraph(style="List Bullet")
    p3.paragraph_format.left_indent = Inches(0.4)
    p3.add_run("1:N")

    add_heading(doc, "Er-Şeması", level=2)
    if ER_IMAGE.exists():
        doc.add_picture(str(ER_IMAGE), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Veritabanı Şeması Görseli Yüklenemedi - Dosya Bulunamadı]")

    add_heading(doc, "İlişkisel (Mantıksal) Şema", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Kategoriler = {")
    run = p.add_run("Kategori ID")
    run.underline = True
    p.add_run(", Kategori Adı, Açıklama}")
    
    p = doc.add_paragraph()
    p.add_run("Ürünler = {")
    run = p.add_run("Ürün ID")
    run.underline = True
    p.add_run(", +Kategori ID, Ürün Adı, Marka, Birim Fiyat, Stok Miktarı, Barkod}")
    
    p = doc.add_paragraph()
    p.add_run("Müşteriler = {")
    run = p.add_run("Müşteri ID")
    run.underline = True
    p.add_run(", Ad, Soyad, Telefon, E-posta, Adres}")
    
    p = doc.add_paragraph()
    p.add_run("Satışlar = {")
    run = p.add_run("Satış ID")
    run.underline = True
    p.add_run(", +Ürün ID, +Müşteri ID, Adet, Birim Fiyat, Toplam Tutar, Satış Tarihi}")
    doc.add_paragraph()

    sql_text = (ROOT / "sql" / "laleli_kozmetik_database.sql").read_text(encoding="utf-8")
    add_heading(doc, "ADIM-3: Veritabanı Programlama")
    doc.add_paragraph(
        "Veritabanı MySQL üzerinde hazırlanmıştır. CRUD işlemleri stored procedure ile yapılır. "
        "Satış eklenmeden önce trigger stok miktarını kontrol eder ve yeterli stok varsa ürün stok miktarını düşürür. "
        "KDV hesaplaması için fonksiyon kullanılmıştır."
    )
    add_code(doc, "Trigger: Satış Yapılınca Stok Düşürme", sql_text[sql_text.index("CREATE TRIGGER"):sql_text.index("CREATE PROCEDURE sp_kategori_ekle")])
    add_code(doc, "Function: KDV'li Fiyat", sql_text[sql_text.index("CREATE FUNCTION"):sql_text.index("CREATE TRIGGER")])
    add_code(doc, "Örnek Stored Procedure: Ürün Ekleme", sql_text[sql_text.index("CREATE PROCEDURE sp_urun_ekle"):sql_text.index("CREATE PROCEDURE sp_urun_guncelle")])

    add_heading(doc, "ADIM-4: Uygulama Geliştirme")
    doc.add_paragraph(
        "C# uygulaması N-katmanlı mimari ile hazırlanmıştır. Data Access Layer yalnızca MySQL bağlantısı ve "
        "stored procedure çağrılarını içerir. Business Logic Layer alan doğrulama ve stok kontrolü yapar. "
        "Presentation Layer Windows Forms ekranlarını içerir."
    )
    add_table(
        doc,
        ["Katman", "Proje/Klasör", "Görev"],
        [
            ["DAL", "LaleliKozmetik.DAL", "MySqlConnection, MySqlCommand ve SP çağrıları"],
            ["BLL", "LaleliKozmetik.BLL", "Ürün, müşteri ve satış iş kuralları"],
            ["UI", "LaleliKozmetik.UI", "Kategoriler, Ürün Yönetimi, Müşteri Kaydı ve Satış Ekranı"],
        ],
    )

    add_code(doc, "BLL Stok Kontrolü", (ROOT / "LaleliKozmetik.BLL" / "SaleService.cs").read_text(encoding="utf-8"))
    add_code(doc, "DAL Satış SP Çağrısı", (ROOT / "LaleliKozmetik.DAL" / "SaleRepository.cs").read_text(encoding="utf-8"))

    add_heading(doc, "Ekranlar")
    add_bullets(
        doc,
        [
            "Kategoriler: Kategori ekleme, silme, güncelleme ve listeleme işlemleri yapılır.",
            "Ürün Yönetimi: Ürün ekleme, silme, güncelleme ve listeleme işlemleri yapılır.",
            "Müşteri Kaydı: Müşteri ekleme, silme, güncelleme ve listeleme işlemleri yapılır.",
            "Satış Ekranı: Satış ekleme, silme, güncelleme ve listeleme işlemleri yapılır; satış sonrası stok miktarı otomatik düşer.",
        ],
    )
    for image, caption in [
        (CATEGORY_SCREEN, "Kategori Yönetimi Ekranı"),
        (PRODUCT_SCREEN, "Ürün Yönetimi Ekranı"),
        (CUSTOMER_SCREEN, "Müşteri Kaydı Ekranı"),
        (SALE_SCREEN, "Satış Ekranı"),
    ]:
        if image.exists():
            doc.add_paragraph(caption).runs[0].bold = True
            doc.add_picture(str(image), width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[{caption} Görseli Yüklenemedi - Dosya Bulunamadı]")

    add_heading(doc, "ADIM-5: Teslimat ve Video Akışı")
    add_bullets(
        doc,
        [
            "SQL scripti: sql/laleli_kozmetik_database.sql dosyası MySQL üzerinde çalıştırılır.",
            "C# proje: LaleliKozmetik.sln Visual Studio ile açılır ve çalıştırılır.",
            "Video 1: MySQL tarafında trigger ve function gösterilir.",
            "Video 2: Uygulamada kategori, ürün, müşteri ve satış CRUD işlemleri gösterilir.",
            "Video 3: Satış yapıldıktan sonra stok miktarının düştüğü, satış silinince stok iadesi olduğu gösterilir.",
            "GitHub: Tüm proje, SQL scripti, ER diyagramı ve rapor yüklenir.",
        ],
    )

    doc.add_paragraph()
    p_video = doc.add_paragraph()
    p_video.add_run("Video Sunum Linki: ").bold = True
    p_video.add_run("__________________________________________________")
    
    doc.add_paragraph()
    p_github = doc.add_paragraph()
    p_github.add_run("GitHub Proje Deposu (Repository) Linki: ").bold = True
    p_github.add_run("https://github.com/emirllli/LaleliKozmetik-Veritabani-Projesi")

    # Try writing to standard output path
    try:
        doc.save(OUTPUT)
        print(f"SUCCESS: {OUTPUT}")
    except PermissionError:
        # Fallback to alternate path if locked by Word
        alt_output = DOCS / "Laleli_Kozmetik_Final_Rapor_GUNCEL.docx"
        doc.save(alt_output)
        print(f"SUCCESS (Alternate due to lock): {alt_output}")


if __name__ == "__main__":
    main()
